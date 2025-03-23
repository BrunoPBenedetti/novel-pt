import os
import tempfile
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Callable
from datetime import datetime
from .web_scraper import WebScraper
from .translator import Translator
from docx import Document

class ChapterManager:
    def __init__(self, novel_data: Dict, progress_callback: Optional[Callable[[int, str], None]] = None, config: Optional['Config'] = None):
        """Inicializa o gerenciador de capítulos."""
        self.novel_data = novel_data
        self.scraper = WebScraper()
        self.translator = Translator()
        self.progress_callback = progress_callback or (lambda x, y: None)
        self.config = config

        # Cria diretórios temporários
        self.temp_dir = Path(tempfile.mkdtemp(prefix="novel_pt_"))
        self.raw_dir = self.temp_dir / "raw"
        self.translated_dir = self.temp_dir / "translated"
        self.raw_dir.mkdir()
        self.translated_dir.mkdir()

        self.log("Iniciando processamento de capítulos...")
        self.log(f"Diretório temporário: {self.temp_dir}")

    def log(self, message: str, progress: int = 0):
        """Registra uma mensagem e atualiza o progresso."""
        print(f"[{datetime.now().strftime('%H:%M:%S')}] {message}")
        self.progress_callback(progress, message)

    def download_chapters(self, start_chapter: int, end_chapter: int) -> bool:
        """Baixa os capítulos da novel."""
        try:
            current_chapter = start_chapter
            total_chapters = end_chapter - start_chapter
            current_url = self.novel_data['current_url']

            self.log(f"Baixando {total_chapters} capítulos...")
            self.log(f"Capítulo inicial: {start_chapter}")
            self.log(f"Capítulo final: {end_chapter}")

            while current_chapter <= end_chapter:
                self.log(f"Baixando capítulo {current_chapter}...")
                self.log(f"URL: {current_url}")

                try:
                    # Obtém o conteúdo da página
                    content = self.scraper.get_page(current_url)
                    if not content:
                        self.log(f"❌ Falha ao obter conteúdo do capítulo {current_chapter}")
                        return False

                    # Extrai o texto do capítulo usando o xpath da novel
                    text = self.scraper.extract_text(content, self.novel_data['content_xpath'])
                    if not text:
                        self.log(f"❌ Falha ao extrair texto do capítulo {current_chapter}")
                        return False

                    # Salva o capítulo
                    chapter_file = self.raw_dir / f"chapter_{current_chapter}.txt"
                    chapter_file.write_text(text, encoding='utf-8')
                    self.log(f"✅ Capítulo {current_chapter} salvo em: {chapter_file}")

                    # Atualiza o progresso
                    progress = int((current_chapter - start_chapter) / total_chapters * 100)
                    self.progress_callback(progress, f"Baixando capítulo {current_chapter}...")

                    # Encontra a URL do próximo capítulo
                    next_url = self.scraper.find_next_chapter_url(self.novel_data['next_chapter_xpath'])
                    if next_url:
                        self.log(f"Próximo capítulo encontrado: {next_url}")
                        current_url = next_url  # Atualiza a URL para o próximo capítulo
                    else:
                        self.log("⚠️ Não foi possível encontrar o próximo capítulo")
                        if current_chapter < end_chapter:
                            self.log("❌ Não é possível continuar sem a URL do próximo capítulo")
                            return False

                    current_chapter += 1

                except Exception as e:
                    self.log(f"❌ Erro ao processar capítulo {current_chapter}: {str(e)}")
                    return False

            self.log("✅ Todos os capítulos foram baixados com sucesso!")
            return True

        except Exception as e:
            self.log(f"❌ Erro ao baixar capítulos: {str(e)}")
            return False

    def translate_chapters(self) -> bool:
        """Traduz os capítulos baixados."""
        try:
            # Lista todos os arquivos de capítulos
            chapter_files = list(self.raw_dir.glob("chapter_*.txt"))
            total_chapters = len(chapter_files)

            if total_chapters == 0:
                self.log("❌ Nenhum capítulo encontrado para traduzir")
                return False

            self.log(f"Traduzindo {total_chapters} capítulos...")

            # Traduz cada capítulo
            for i, chapter_file in enumerate(chapter_files, 1):
                try:
                    self.log(f"Traduzindo capítulo {i}/{total_chapters}...")

                    # Lê o conteúdo do capítulo
                    content = chapter_file.read_text(encoding='utf-8')
                    if not content.strip():
                        self.log(f"⚠️ Capítulo {i} está vazio, pulando...")
                        continue

                    # Traduz o conteúdo
                    translated_content = self.translator.translate_text(content)
                    if not translated_content:
                        self.log(f"❌ Falha ao traduzir capítulo {i}")
                        return False

                    # Salva o capítulo traduzido
                    translated_file = self.translated_dir / f"chapter_{i}.txt"
                    translated_file.write_text(translated_content, encoding='utf-8')
                    self.log(f"✅ Capítulo {i} traduzido e salvo")

                    # Atualiza o progresso
                    progress = int((i / total_chapters) * 100)
                    self.progress_callback(progress, f"Traduzindo capítulo {i}/{total_chapters}...")

                except Exception as e:
                    self.log(f"❌ Erro ao traduzir capítulo {i}: {str(e)}")
                    return False

            self.log("✅ Todos os capítulos foram traduzidos com sucesso!")
            return True

        except Exception as e:
            self.log(f"❌ Erro ao traduzir capítulos: {str(e)}")
            return False

    def merge_chapters(self) -> Optional[str]:
        """Combina os capítulos traduzidos em um único arquivo."""
        try:
            # Lista todos os arquivos de capítulos traduzidos
            chapter_files = sorted(self.translated_dir.glob("chapter_*.txt"))
            if not chapter_files:
                self.log("❌ Nenhum capítulo traduzido encontrado")
                return None

            # Cria o diretório de saída se não existir
            output_dir = Path(self.novel_data['output_dir'])
            output_dir.mkdir(parents=True, exist_ok=True)

            # Nome do arquivo de saída
            novel_name = self.novel_data['name']
            output_format = self.novel_data.get('format', 'DOCX')
            output_file = output_dir / f"{novel_name}.{output_format.lower()}"

            if output_format == 'DOCX':
                # Cria um novo documento DOCX
                doc = Document()

                # Adiciona os capítulos ao documento
                for i, chapter_file in enumerate(chapter_files, 1):
                    try:
                        # Lê o conteúdo do capítulo
                        content = chapter_file.read_text(encoding='utf-8')

                        # Adiciona o número do capítulo se necessário
                        if self.novel_data.get('show_chapter_number', True):
                            doc.add_paragraph(f"\nCapítulo {i}\n", style='Heading 1')
                            doc.add_paragraph()  # Espaço após o título

                        # Adiciona o conteúdo
                        doc.add_paragraph(content)
                        doc.add_paragraph()  # Espaço entre capítulos

                        self.log(f"✅ Capítulo {i} adicionado ao arquivo final")

                    except Exception as e:
                        self.log(f"❌ Erro ao processar capítulo {i}: {str(e)}")
                        return None

                # Salva o documento
                doc.save(str(output_file))
            else:  # TXT
                # Combina os capítulos em um arquivo TXT
                with open(output_file, 'w', encoding='utf-8') as f:
                    for i, chapter_file in enumerate(chapter_files, 1):
                        try:
                            # Lê o conteúdo do capítulo
                            content = chapter_file.read_text(encoding='utf-8')

                            # Adiciona o número do capítulo se necessário
                            if self.novel_data.get('show_chapter_number', True):
                                f.write(f"\nCapítulo {i}\n\n")

                            # Escreve o conteúdo
                            f.write(content)
                            f.write("\n\n")  # Espaço entre capítulos

                            self.log(f"✅ Capítulo {i} adicionado ao arquivo final")

                        except Exception as e:
                            self.log(f"❌ Erro ao processar capítulo {i}: {str(e)}")
                            return None

            self.log(f"✅ Arquivo final gerado com sucesso: {output_file}")
            return str(output_file)

        except Exception as e:
            self.log(f"❌ Erro ao gerar arquivo final: {str(e)}")
            return None

    def cleanup(self):
        """Remove os arquivos temporários."""
        try:
            if hasattr(self, 'temp_dir') and self.temp_dir.exists():
                self.log("🧹 Limpando arquivos temporários...", 95)
                shutil.rmtree(self.temp_dir)
                self.log("✅ Arquivos temporários removidos com sucesso", 100)
        except Exception as e:
            self.log(f"⚠️ Erro ao limpar arquivos temporários: {str(e)}")

    def process_chapters(self, start_chapter: int, batch_size: int = 1) -> Optional[str]:
        """Processa os capítulos da novel."""
        try:
            # Calcula o número total de capítulos a serem traduzidos
            current_chapter = start_chapter
            end_chapter = start_chapter + batch_size - 1
            total_chapters = batch_size

            self.log(f"Processando {total_chapters} capítulos...")
            self.log(f"Capítulo inicial: {current_chapter}")
            self.log(f"Capítulo final: {end_chapter}")
            self.log(f"Tamanho do lote: {batch_size}")

            # Baixa os capítulos
            if not self.download_chapters(current_chapter, end_chapter):
                self.log("❌ Falha ao baixar os capítulos")
                return None

            # Traduz os capítulos
            if not self.translate_chapters():
                self.log("❌ Falha ao traduzir os capítulos")
                return None

            # Gera o arquivo final
            output_file = self.merge_chapters()
            if not output_file:
                self.log("❌ Falha ao gerar o arquivo final")
                return None

            # Atualiza o capítulo atual e a URL apenas se todo o processo foi bem sucedido
            if self.config:
                # Obtém a URL do próximo capítulo
                next_url = self.scraper.find_next_chapter_url(self.novel_data['next_chapter_xpath'])

                # Cria uma cópia dos dados atuais da novel
                update_data = self.novel_data.copy()

                # Atualiza apenas os campos necessários
                update_data.update({
                    'current_chapter': end_chapter + 1,
                    'current_url': next_url if next_url else self.novel_data['current_url']
                })

                self.config.update_novel(self.novel_data['id'], update_data)
                self.log(f"✅ Capítulo atual atualizado para: {end_chapter + 1}")
                if next_url:
                    self.log(f"✅ URL atual atualizada para: {next_url}")
                else:
                    self.log("⚠️ Não foi possível encontrar a URL do próximo capítulo")

            return output_file

        except Exception as e:
            self.log(f"❌ Erro ao processar capítulos: {str(e)}")
            return None

    def __del__(self):
        """Destrutor para garantir a limpeza dos recursos."""
        try:
            self.cleanup()
        except:
            pass  # Ignora erros no destrutor
